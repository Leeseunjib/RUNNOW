# -*- coding: utf-8 -*-
import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D0D7DE", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_word_report(output_path):
    doc = docx.Document()

    # 페이지 여백 설정 (Narrow)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 기본 스타일 폰트
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x24, 0x29, 0x2F)

    # 표지 / 헤더 타이틀
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_icon = title_p.add_run("⚡ ")
    run_icon.font.size = Pt(22)
    run_title = title_p.add_run("[마스터 종합 보고서] RUNNOW 앱 개발 전수 분석")
    run_title.font.name = 'Malgun Gothic'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("투입 Skill · 핵심 기술 · 활용 언어 · 종합 기획 총망라 정본 (SSOT)")
    run_sub.font.name = 'Malgun Gothic'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # 메타데이터 박스 (표)
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    meta_data = [
        ("문서 코드", "BSC-REP-20260904-RUNNOW-MASTER"),
        ("프로젝트", "RUNNOW (RunGotchi - Nike x Tamagotchi Runner & AI Training OS)"),
        ("작성 일자", "2026년 09월 04일"),
        ("작성자 / 보고 대상", "CTO 거누 (Backend/Infra Lead) ➔ 이건우 대표님 (BeausCreators CEO)"),
        ("문서 상태", "공식 확정 정본 (SSOT) / 전사 지식 자산 등록 완료")
    ]
    set_table_borders(meta_table, color="CBD5E1", sz="6")
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(1.8)
        cell_val.width = Inches(5.0)
        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)

        p0 = cell_lbl.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        p1 = cell_val.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_section_header(number_str, title_str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r_num = h.add_run(f"{number_str}. ")
        r_num.font.name = 'Malgun Gothic'
        r_num.font.size = Pt(14)
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)  # Cyan/Sky

        r_txt = h.add_run(title_str)
        r_txt.font.name = 'Malgun Gothic'
        r_txt.font.size = Pt(14)
        r_txt.font.bold = True
        r_txt.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    def add_bullet(p, bold_prefix, text):
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.2)
        r_b = p.add_run(bold_prefix + ": ")
        r_b.font.bold = True
        r_b.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        r_t = p.add_run(text)
        r_t.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # 1. 프로젝트 개요 및 핵심 비전
    add_section_header("1", "📌 프로젝트 개요 및 핵심 비전 (Executive Summary)")
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(6)
    p_desc.add_run(
        "RUNNOW는 글로벌 러닝 트래커의 표준인 Nike Run Club(NRC)의 세련된 고대비 다크/볼트(#CCFF00) UI와 "
        "다마고치(Tamagotchi)의 사이버 펫 육성·진화 시스템, 그리고 Google MediaPipe 온디바이스 AI 비전 인터랙티브 트레이닝 엔진을 "
        "완벽하게 융합한 차세대 하이퍼-게이미피케이션 피트니스 OS입니다."
    )

    bp1 = doc.add_paragraph(style='List Bullet')
    add_bullet(bp1, "슬로건", "“달리는 만큼 진화하는 나만의 사이버 펫 러닝 OS — Run, Grow, Evolve!”")
    bp2 = doc.add_paragraph(style='List Bullet')
    add_bullet(bp2, "동반 모험", "고독하고 지루한 러닝을 사이버 펫 '볼트'와의 감정적 동반 모험으로 전환")
    bp3 = doc.add_paragraph(style='List Bullet')
    add_bullet(bp3, "즉각 도파민 회로", "체중 감량 등 지연된 보상 대신 수 초 내 펫 레벨업/스탯 상승/코인 보상 지급")
    bp4 = doc.add_paragraph(style='List Bullet')
    add_bullet(bp4, "올인원 실내외", "야외 GPS 러닝과 실내 AI 웹캠 모션 트레이닝(푸시업/스쿼트/윗몸일으키기)의 365일 연동")
    bp5 = doc.add_paragraph(style='List Bullet')
    add_bullet(bp5, "제로 인프라 비용", "WebGPU 및 MediaPipe 온디바이스 처리로 서버 AI 비용 $0 및 100% 프라이버시 달성")

    # 2. 투입된 Skill 전수 분석
    add_section_header("2", "🛠️ 투입된 Skill (에이전트 및 특화 스킬셋 전수)")
    doc.add_paragraph("앱 기획부터 비전 AI, 결제 및 배포 파이프라인 전 과정에 투입된 특화 스킬 목록입니다.")

    skills_data = [
        ("ondevice-ai-domain-routing", "자체 전용 스킬", "카페24(beauscreators.com) 멀티앱 CNAME 라우팅, 무료 SSL 연동 및 WebGPU 기반 15대 온디바이스 AI(Gemma2, Llama3.2, DeepSeek 등) 로컬 구동 스킬"),
        ("notebooklm-automation-pipeline", "사전 연구 스킬", "Google NotebookLM API를 활용한 피트니스 게이미피케이션 학술 연구 자동화 (기저핵 습관 루프, UCL 66일 습관 형성 곡선, Octalysis 프레임워크 리포트 도출)"),
        ("StitchMCP (UI/UX 설계)", "디자인 스킬", "Nike Run Club(NRC) 시그니처 형광 볼트(#CCFF00) + OLED 딥 다크(#08090C) 디자인 시스템 구축 및 5대 탭 UI 계층 구조 명세"),
        ("firebase-mcp-server", "백엔드/인프라", "Firebase Auth(익명/구글/이메일), Cloud Firestore, Cloud Storage 연동, 보안 규칙 및 복합 인덱스, Hosting 자동 배포"),
        ("paypal-mcp-server", "글로벌 결제", "PayPal REST API v2 및 JavaScript SDK Buttons 결제 연동, 20종 인게임 아이템 카탈로그 및 주문 생성/캡처 결제 검증 수립"),
        ("accidental-data-loss-prevention", "데이터 보호", "원본 파일 통째 덮어쓰기 금지, 부분 정밀 수정(replace_file_content), 백업 및 0-Byte 유실 방지 수칙 준수"),
        ("bsc-global-rules", "글로벌 운영", "CHLGD(Context, Harness, Loop, Graph, Dynamic) 워크플로우 준수 및 HQ-Branch 양방향 동기화")
    ]

    t_skill = doc.add_table(rows=len(skills_data)+1, cols=3)
    t_skill.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_skill, color="CBD5E1")
    headers = ["스킬명", "구분", "핵심 역할 및 실제 적용 내역"]
    widths = [Inches(1.8), Inches(1.1), Inches(3.9)]

    for c_idx, h_text in enumerate(headers):
        cell = t_skill.rows[0].cells[c_idx]
        cell.width = widths[c_idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)

    for r_idx, (s_name, s_type, s_desc) in enumerate(skills_data):
        row = t_skill.rows[r_idx+1]
        for c_idx, val in enumerate([s_name, s_type, s_desc]):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 3. 적용된 핵심 기술
    add_section_header("3", "💻 적용된 핵심 기술 (Technologies, Frameworks & APIs)")

    h3_1 = doc.add_paragraph()
    h3_1.paragraph_format.space_before = Pt(8)
    h3_1.paragraph_format.space_after = Pt(2)
    r = h3_1.add_run("3.1. 온디바이스(On-Device) AI & 비전 인터랙티브 엔진")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "Google MediaPipe Pose (WASM/WebGL)", "웹캠 영상에서 33개 신체 관절 3D 랜드마크를 60FPS로 로컬 추출. 서버 비용 $0, 프라이버시 100% 보장.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "관절 벡터 각도 연산", "푸시업(어깨-팔꿈치-손목 90°), 스쿼트(골반-무릎-발목 90°), 윗몸일으키기(어깨-골반-무릎 45° 수축) 정밀 판정 머신 구현.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "Web Audio API Synthesizer", "외부 음원 다운로드 없이 브라우저 내장 AudioContext 오스킬레이터로 카운트다운, 성공 차임, 경고음 합성 생성.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "YouTube Iframe API 듀얼 뷰", "상단 50% 트레이너 가이드 영상 + 하단 50% 실시간 웹캠 스켈레톤 라인 오버레이 뷰 분할 배치.")

    h3_2 = doc.add_paragraph()
    h3_2.paragraph_format.space_before = Pt(8)
    h3_2.paragraph_format.space_after = Pt(2)
    r = h3_2.add_run("3.2. 고정밀 실시간 GPS 엔지니어링")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "하버사인(Haversine) 구면 공식", "지구 반경 6,371km 기준 1초 단위 미세 위경도 변화량을 누적 합산하여 출발점 복귀 코스에서도 거리 100% 무손실 보존.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "정수 미터(m) 실시간 표기", "둔감한 0.01km 단위 대신 12m, 154m, 1,250m로 1m 단위 즉각 증가 표기하여 전진 모멘텀 극대화.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "지터 & 안티치트 필터", "정확도 25m 초과 불량 신호 배제, 정지 상태 미세 흔들림 제거, 시속 30km/h 초과 차량 탑승 구간 자동 감지 및 거리 배제.")

    h3_3 = doc.add_paragraph()
    h3_3.paragraph_format.space_before = Pt(8)
    h3_3.paragraph_format.space_after = Pt(2)
    r = h3_3.add_run("3.3. 백엔드 및 결제·PWA 인프라")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "Firebase Authentication", "게스트(익명) 원클릭 시작 지원으로 온보딩 허들 제거 후 Google/이메일 정식 계정 링크 지원.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "Cloud Firestore & Security Rules", "다마고치 상태, 운동 기록, 21일 챌린지, 결제 내역의 실시간 양방향 동기화 및 샌드박스 접근 제어.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "PayPal JavaScript SDK v2", "PayPal Buttons 결제 UI 및 클라이언트 Order Create / Capture 결제 검증 파이프라인.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "PWA (Service Worker & Manifest)", "오프라인 캐싱, 스마트폰 홈 화면 추가(A2HS), 풀스크린 네이티브 앱 UX 구현.")

    # 4. 활용된 언어 및 데이터 포맷
    add_section_header("4", "🔤 활용된 언어 및 데이터 포맷 (Languages & Formats)")
    lang_data = [
        ("JavaScript (ES6+)", "클라이언트 코어 로직", "모듈형 바닐라 JS 객체 지향 및 함수형 설계. 비전 AI, GPS 수학 연산, 상태 머신, 결제 브릿지 구동"),
        ("HTML5", "화면 구조 및 쉘", "index.html(메인 러닝 OS 5대 탭) 및 workout.html(카메라 모션 트레이닝 분할 뷰)의 시맨틱 웹 구성"),
        ("CSS3", "디자인 시스템 & 모션", "CSS 변수 기반 테마 토큰, Flexbox/Grid, 형광 볼트 글래스모피즘, 60FPS 펄스/샤인 애니메이션"),
        ("JSON", "데이터 규격 & 설정", "PWA manifest.json, firebase.json, firestore.indexes.json 및 20종 상점 아이템 카탈로그"),
        ("Firestore Rules", "데이터베이스 보안", "유저별 격리 및 데이터 스키마 유효성 검증 규칙"),
        ("Markdown / Mermaid", "기획 & 아키텍처", "PRD, 시스템 사양서, 플로우차트, 시퀀스 다이어그램 기술")
    ]

    t_lang = doc.add_table(rows=len(lang_data)+1, cols=3)
    t_lang.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_lang, color="CBD5E1")
    l_headers = ["언어 / 포맷", "활용 영역", "상세 설명"]
    l_widths = [Inches(1.8), Inches(1.4), Inches(3.6)]

    for c_idx, h_text in enumerate(l_headers):
        cell = t_lang.rows[0].cells[c_idx]
        cell.width = l_widths[c_idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)

    for r_idx, (l_name, l_type, l_desc) in enumerate(lang_data):
        row = t_lang.rows[r_idx+1]
        for c_idx, val in enumerate([l_name, l_type, l_desc]):
            cell = row.cells[c_idx]
            cell.width = l_widths[c_idx]
            bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5. 종합 기획 내용
    add_section_header("5", "🧠 종합 기획 내용 (Product Planning & Behavioral Mechanics)")

    h5_1 = doc.add_paragraph()
    h5_1.paragraph_format.space_before = Pt(8)
    h5_1.paragraph_format.space_after = Pt(2)
    r = h5_1.add_run("5.1. 행동 심리학 & 옥탈리시스(Octalysis) 게이미피케이션 매트릭스")
    r.font.bold = True
    r.font.size = Pt(11)

    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "CD 1 (거대한 사명감)", "나의 달리기와 땀이 멸종 위기의 사이버 펫 '볼트'를 부화시키고 진화시킨다는 스토리텔링 부여.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "CD 2 (진보와 성취)", "5km, 20km, 50km, 100km 돌파 시 5단계 변신 및 마일스톤 배지 수여로 성취감 가시화.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "CD 3 (창의성 부여)", "20종 장비(러닝화, 바이저, 오라 등) 커스텀 조합으로 나만의 펫 스타일링.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "CD 4 (소유권과 자산)", "달려서 채굴한 볼트코인(VC)으로 상점 아이템 구매 및 희귀 스킨 영구 소장.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "CD 8 (손실 회피 완충)", "운동을 쉬어도 즉시 펫이 사망하지 않고 에너지 저하 상태로 유예 후 스트릭 실드로 회복 유도.")

    h5_2 = doc.add_paragraph()
    h5_2.paragraph_format.space_before = Pt(8)
    h5_2.paragraph_format.space_after = Pt(2)
    r = h5_2.add_run("5.2. 다마고치 5단계 진화 & 3대 스탯 알고리즘")
    r.font.bold = True
    r.font.size = Pt(11)

    evo_data = [
        ("Phase 1", "Runner Egg (러너 에그)", "0.0 km (가입 즉시)", "진동하는 네온 볼트 알", "튜토리얼 완주 시 즉시 부화"),
        ("Phase 2", "Rookie Chick (루키 병아리)", "5.0 km", "볼트 헤드밴드 장착 날쌘 병아리", "일일 러닝 XP +10% 추가 획득"),
        ("Phase 3", "Urban Runner (어반 러너)", "20.0 km", "바람막이를 입은 사이버 여우", "러닝 중 배고픔 소모 -15% 완화"),
        ("Phase 4", "Marathon Master (마라톤 마스터)", "50.0 km", "사이버 글래스/카본화의 흑표범", "칼로리당 볼트코인 획득량 +20%"),
        ("Phase 5", "Cyber Speedster (사이버 신수)", "100.0 km", "번개 오라를 두른 궁극의 신수", "볼트 상점 전 품목 15% 상시 할인")
    ]

    t_evo = doc.add_table(rows=len(evo_data)+1, cols=5)
    t_evo.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_evo, color="CBD5E1")
    e_headers = ["단계", "캐릭터명", "진화 조건", "외형 컨셉", "패시브 버프"]
    e_widths = [Inches(0.9), Inches(1.6), Inches(1.0), Inches(1.7), Inches(1.6)]

    for c_idx, h_text in enumerate(e_headers):
        cell = t_evo.rows[0].cells[c_idx]
        cell.width = e_widths[c_idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(8.5)

    for r_idx, row_vals in enumerate(evo_data):
        row = t_evo.rows[r_idx+1]
        for c_idx, val in enumerate(row_vals):
            cell = row.cells[c_idx]
            cell.width = e_widths[c_idx]
            bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if c_idx in [0, 1]:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    h5_3 = doc.add_paragraph()
    h5_3.paragraph_format.space_before = Pt(8)
    h5_3.paragraph_format.space_after = Pt(2)
    r = h5_3.add_run("5.3. 비즈니스 모델(BM) & 20종 상점 상품 카탈로그")
    r.font.bold = True
    r.font.size = Pt(11)

    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "러닝화 (4종)", "Volt Pegasus Turbo ($4.99), Cyber VaporFly Next% ($9.99), Carbon Streak X ($14.99), Alpha Aero Fly ($19.99)")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "에너지 (4종)", "Volt Hydration ($0.99), Nano Electrolyte Gel ($1.99), Beast Protein Shake ($2.99), Phoenix Elixir ($4.99)")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "스킨/오라 (4종)", "Cyberpunk Neon Visor ($2.99), Night Tracksuit ($5.99), Golden Champion Aura ($8.99), Midnight Ninja Hoodie ($6.99)")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "웨어러블 (4종)", "Titanium GPS Pro Watch ($3.99), Aero Speed Sunglasses ($2.49), Reflex LED Armband ($1.99), SoundPulse Headband ($3.49)")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "프리미엄 패스 (4종)", "3-Week Double XP Pass ($9.99), Daily Streak Shield ($4.99), Master Trophy Box ($12.99), VIP Diamond Club ($19.99/월)")

    h5_4 = doc.add_paragraph()
    h5_4.paragraph_format.space_before = Pt(8)
    h5_4.paragraph_format.space_after = Pt(2)
    r = h5_4.add_run("5.4. 온디바이스 15대 AI 라인업 및 카페24 CNAME 라우팅")
    r.font.bold = True
    r.font.size = Pt(11)

    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "카페24 CNAME 라우팅", "쇼핑몰(beauscreators.com) 중단 없이 runnow.beauscreators.com, nutrios.beauscreators.com 등으로 서브도메인 확장.")
    bp = doc.add_paragraph(style='List Bullet')
    add_bullet(bp, "온디바이스 15대 AI", "초경량 실시간 코칭 7종(Gemma2, Llama3.2 등), 비전 멀티모달 4종(PaliGemma2, Phi-3.5 Vision 등), 심층 추론 2종(DeepSeek-R1), 음성 인식 2종(Whisper Web).")

    # 6. 핵심 파일별 매핑 요약
    add_section_header("6", "📂 핵심 파일별 역할 및 시스템 매핑")
    file_map = [
        ("index.html", "메인 러닝 OS 앱 쉘 (5대 탭 UI, NRC Live HUD, 다마고치 룸)"),
        ("workout.html", "AI 웹캠 모션 트레이닝 화면 (MediaPipe + YouTube 듀얼 뷰)"),
        ("styles.css", "NRC 시그니처 볼트/다크 테마 디자인 시스템 (글래스모피즘)"),
        ("app.js", "전체 앱 라이프사이클 및 화면 라우팅, 전역 상태 통합 관리"),
        ("motionTracker.js", "MediaPipe 관절 각도 벡터 계산 및 실시간 운동 자세 판정 머신"),
        ("motionSound.js", "Web Audio API 기반 비프음, 차임벨, 버저 사운드 신디사이저"),
        ("gpsRunner.js", "고정밀 GPS 엔진, 하버사인 거리 누적, 지터/속도 필터"),
        ("tamagotchi.js", "5단계 진화 트리, 3대 스탯 시뮬레이션, 먹이/훈련 인터랙션"),
        ("challenge.js", "21일 습관 부트캠프, 일일 미션 체크, 신체 지수(BMR/BMI) 관리"),
        ("quests.js", "일일 퀘스트, 마일스톤 업적 달성 및 보상 엔진"),
        ("catalog.js", "20종 인게임 상점 아이템 정본 데이터 모델"),
        ("paypalBridge.js", "PayPal JS SDK 주문 생성 및 캡처, 인벤토리 지급 브릿지"),
        ("firebaseClient.js", "Firebase Auth 및 Cloud Firestore 실시간 동기화 클라이언트"),
        ("firebaseConfig.js", "Firebase 프로젝트 키 및 환경 구성"),
        ("firestore.rules", "데이터 보안 및 유저별 접근 제어 규칙"),
        ("sw.js", "PWA Service Worker (오프라인 캐싱 및 리소스 프리패치)"),
        ("manifest.json", "PWA 웹앱 매니페스트 (홈 화면 설치 및 아이콘 정의)"),
        ("RUNNOW_Master_PRD.md", "BeausCreators 전사 마스터 기획서 (SSOT)")
    ]

    t_file = doc.add_table(rows=len(file_map)+1, cols=2)
    t_file.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_file, color="CBD5E1")
    f_headers = ["파일명", "역할 및 담당 시스템"]
    f_widths = [Inches(2.2), Inches(4.6)]

    for c_idx, h_text in enumerate(f_headers):
        cell = t_file.rows[0].cells[c_idx]
        cell.width = f_widths[c_idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)

    for r_idx, (fname, frole) in enumerate(file_map):
        row = t_file.rows[r_idx+1]
        for c_idx, val in enumerate([fname, frole]):
            cell = row.cells[c_idx]
            cell.width = f_widths[c_idx]
            bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    # 저장
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated docx: {output_path}")

if __name__ == "__main__":
    local_target = r"c:\BeausCreators\03.Research\바이브코딩 연구\Test_proj\proj_01\docs\2026-09-04_RUNNOW_Tech_and_Planning_Master_Report.docx"
    build_word_report(local_target)
